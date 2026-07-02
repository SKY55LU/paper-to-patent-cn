#!/usr/bin/env python3
# Copyright (c) 2026 LYJ. All rights reserved.
"""Generate a final-format Chinese patent DOCX and optional review PDF from JSON."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Iterable, Mapping, Sequence

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ImportError as exc:  # pragma: no cover
    raise SystemExit("Missing dependency: python-docx. Install it in the active Python environment.") from exc


MAIN_HEADINGS = [
    "说   明   书   摘   要",
    "摘    要   附   图",
    "权  利  要  求  书",
    "说   明   书",
    "说 明 书 附 图",
]


def strip_publication_number(text: str) -> str:
    return re.sub(r"^\s*(?:\[\d{4}\]|\d{4})\s*", "", text or "").strip()


def as_paragraphs(value: object) -> list[str]:
    if value is None:
        return []
    if isinstance(value, str):
        return [strip_publication_number(p) for p in re.split(r"\n+", value) if strip_publication_number(p)]
    if isinstance(value, Sequence) and not isinstance(value, (bytes, bytearray)):
        return [strip_publication_number(str(p)) for p in value if strip_publication_number(str(p))]
    return [strip_publication_number(str(value))]


def set_run_font(run, size_pt: float = 12, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size_pt)
    if bold:
        run.bold = True


def set_paragraph_line_spacing(paragraph, points: int = 22) -> None:
    paragraph.paragraph_format.line_spacing = Pt(points)


def clear_paragraph_borders(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is not None:
        p_pr.remove(p_bdr)


def add_bottom_border(paragraph) -> None:
    clear_paragraph_borders(paragraph)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")


def add_main_heading(doc: Document, text: str, page_break_before: bool = False) -> None:
    p = doc.add_paragraph()
    clear_paragraph_borders(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_line_spacing(p)
    if page_break_before:
        p.paragraph_format.page_break_before = True
    add_bottom_border(p)
    run = p.add_run(text)
    set_run_font(run, 12, False)


def add_title(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    clear_paragraph_borders(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_line_spacing(p)
    run = p.add_run(title)
    set_run_font(run, 12, True)


def add_cover_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    clear_paragraph_borders(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_line_spacing(p)
    p.paragraph_format.first_line_indent = None
    run = p.add_run(text)
    set_run_font(run, 12, False)


def add_cover_page(doc: Document, title: str, data: Mapping[str, object]) -> None:
    cover = data.get("cover_fields")
    if isinstance(cover, Mapping):
        values = {
            "专利名称": str(cover.get("专利名称") or title),
            "校内识别号": str(cover.get("校内识别号") or "Z"),
            "专利类型": str(cover.get("专利类型") or "国内专利"),
            "专利类别": str(cover.get("专利类别") or "发明专利"),
            "申请人": str(cover.get("申请人") or ""),
            "发明人": str(cover.get("发明人") or ""),
            "第一发明人身份证号": str(cover.get("第一发明人身份证号") or ""),
            "申报人": str(cover.get("申报人") or ""),
            "申报人电话": str(cover.get("申报人电话") or ""),
            "申报人邮箱": str(cover.get("申报人邮箱") or ""),
            "联系人": str(cover.get("联系人") or ""),
            "联系人电话": str(cover.get("联系人电话") or ""),
            "联系人邮箱": str(cover.get("联系人邮箱") or ""),
        }
    else:
        values = {
            "专利名称": title,
            "校内识别号": "Z",
            "专利类型": "国内专利",
            "专利类别": "发明专利",
            "申请人": "",
            "发明人": "",
            "第一发明人身份证号": "",
            "申报人": "",
            "申报人电话": "",
            "申报人邮箱": "",
            "联系人": "",
            "联系人电话": "",
            "联系人邮箱": "",
        }
    for key, value in values.items():
        add_cover_line(doc, f"{key}：{value}")


def add_subheading(doc: Document, text: str, page_break_before: bool = False) -> None:
    p = doc.add_paragraph()
    clear_paragraph_borders(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_line_spacing(p)
    if page_break_before:
        p.paragraph_format.page_break_before = True
    run = p.add_run(text)
    set_run_font(run, 12, True)


def add_body_paragraph(doc: Document, text: str, indent: bool = True) -> None:
    clean = strip_publication_number(text)
    if not clean:
        return
    p = doc.add_paragraph()
    clear_paragraph_borders(p)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_line_spacing(p)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(clean)
    set_run_font(run, 12, False)


def add_claim(doc: Document, number: int, text: str) -> None:
    clean = strip_publication_number(text)
    clean = re.sub(r"^\s*\d+[\.、]\s*", "", clean)
    if not clean.endswith("。"):
        clean += "。"
    add_body_paragraph(doc, f"{number}. {clean}", indent=False)


def add_image(doc: Document, image_path: Path, caption: str | None = None, width_cm: float = 13.0) -> None:
    if not image_path.exists():
        raise FileNotFoundError(f"Drawing asset not found: {image_path}")
    p = doc.add_paragraph()
    clear_paragraph_borders(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph()
        clear_paragraph_borders(cap)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_line_spacing(cap)
        r = cap.add_run(caption)
        set_run_font(r, 12, False)


def add_figure_caption(doc: Document, caption: str | None) -> None:
    if not caption or not caption.strip():
        return
    p = doc.add_paragraph()
    clear_paragraph_borders(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_line_spacing(p)
    run = p.add_run(caption.strip())
    set_run_font(run, 12, False)


def resolve_path(base: Path, value: str) -> Path:
    path = Path(value)
    return path if path.is_absolute() else base / path


def build_docx(data: Mapping[str, object], json_path: Path, output_docx: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.7)

    title = str(data.get("title") or "").strip()
    if not title:
        raise ValueError("JSON field 'title' is required.")

    abstract = as_paragraphs(data.get("abstract"))
    claims = as_paragraphs(data.get("claims"))
    if not claims:
        raise ValueError("JSON field 'claims' must contain at least one claim.")

    description = data.get("description") or {}
    if not isinstance(description, Mapping):
        raise ValueError("JSON field 'description' must be an object.")

    drawings = data.get("drawing_assets") or []
    if drawings and not isinstance(drawings, Sequence):
        raise ValueError("JSON field 'drawing_assets' must be a list.")

    add_cover_page(doc, title, data)

    add_main_heading(doc, MAIN_HEADINGS[0], page_break_before=True)
    for para in abstract:
        add_body_paragraph(doc, para)

    add_main_heading(doc, MAIN_HEADINGS[1], page_break_before=True)
    if drawings and isinstance(drawings[0], Mapping):
        add_figure_caption(doc, str(drawings[0].get("caption") or "图1"))

    add_main_heading(doc, MAIN_HEADINGS[2], page_break_before=True)
    for i, claim in enumerate(claims, 1):
        add_claim(doc, i, claim)

    add_main_heading(doc, MAIN_HEADINGS[3], page_break_before=True)
    add_title(doc, title)

    section_map = [
        ("技术领域", "technical_field"),
        ("背景技术", "background"),
        ("发明内容", "summary"),
        ("附图说明", "drawing_description"),
        ("具体实施方式", "embodiments"),
    ]
    for heading, key in section_map:
        add_subheading(doc, heading)
        for para in as_paragraphs(description.get(key)):
            add_body_paragraph(doc, para)
        if key == "summary":
            for para in as_paragraphs(description.get("benefits")):
                add_body_paragraph(doc, para)

    add_main_heading(doc, MAIN_HEADINGS[4], page_break_before=True)
    for item in drawings:
        if isinstance(item, Mapping):
            add_figure_caption(doc, str(item.get("caption") or "").strip() or None)
        else:
            add_figure_caption(doc, str(item).strip() or None)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)


def make_review_pdf(docx_path: Path, pdf_path: Path) -> None:
    try:
        from PIL import Image, ImageDraw, ImageFont
        from docx import Document as ReadDocument
    except ImportError as exc:  # pragma: no cover
        raise SystemExit("Missing dependency for PDF review copy: pillow and python-docx.") from exc

    document = ReadDocument(docx_path)
    text_lines: list[str] = []
    for para in document.paragraphs:
        if para.text.strip():
            text_lines.append(para.text.strip())

    page_w, page_h = 1240, 1754
    margin = 110
    line_h = 34
    try:
        font = ImageFont.truetype("simsun.ttc", 22)
    except OSError:
        font = ImageFont.load_default()

    pages: list[Image.Image] = []
    image = Image.new("RGB", (page_w, page_h), "white")
    draw = ImageDraw.Draw(image)
    y = margin

    def wrap(line: str, width: int = 46) -> Iterable[str]:
        for i in range(0, len(line), width):
            yield line[i : i + width]

    for line in text_lines:
        for part in wrap(line):
            if y > page_h - margin:
                pages.append(image)
                image = Image.new("RGB", (page_w, page_h), "white")
                draw = ImageDraw.Draw(image)
                y = margin
            draw.text((margin, y), part, fill="black", font=font)
            y += line_h
        y += 10
    pages.append(image)
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    pages[0].save(pdf_path, save_all=True, append_images=pages[1:])


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("json", type=Path, help="Structured patent JSON.")
    parser.add_argument("--output-dir", type=Path, default=Path("outputs_patent"))
    parser.add_argument("--basename", default=None)
    parser.add_argument("--pdf", action="store_true", help="Also create a simple image-based PDF review copy.")
    args = parser.parse_args()

    data = json.loads(args.json.read_text(encoding="utf-8"))
    basename = args.basename or re.sub(r'[\\/:*?"<>|]+', "_", str(data.get("title") or "patent_application"))
    output_docx = args.output_dir / f"{basename}.docx"
    output_pdf = args.output_dir / f"{basename}.pdf"

    build_docx(data, args.json.resolve(), output_docx)
    print(f"DOCX: {output_docx}")
    if args.pdf:
        make_review_pdf(output_docx, output_pdf)
        print(f"PDF: {output_pdf}")


if __name__ == "__main__":
    main()



